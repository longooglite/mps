MY_API_KEY = '392bff629e6d07593968592bbb629b0f'



#http://api.elsevier.com:80/content/article/scopus_id/0037070197?apiKey=392bff629e6d07593968592bbb629b0f

#http://api.elsevier.com/content/search/index:SCOPUS?query=[query]

#http://www.scopus.com/results/affiliationResults.uri?selectionPageSearch=afsp&affilName=University+of+Michigan&cl=t&sort=afcnt-f&src=s&sid=72B2D67C88B703BB6B7E888764A26CFC.kqQeWtawXauCyC8ghhRGJg%3a230&sot=afsp&sdt=aff&sl=29&s=AFFIL%28University+of+Michigan%29&origin=SearchAffiliationLookup&txGid=0

#http://www.scopus.com/results/authorNamesList.uri?origin=searchauthorlookup&src=al&edit=&poppUp=&basicTab=&affiliationTab=&advancedTab=&st1=Chung&st2=kevin&institute=University+of+Michigan&_exactSearch=on&orcidId=&authSubject=LFSC&_authSubject=on&authSubject=HLSC&_authSubject=on&authSubject=PHSC&_authSubject=on&authSubject=SOSC&_authSubject=on&s=AUTH-LAST-NAME%28Chung%29+AND+AUTH--FIRST%28kevin%29&sdt=al&sot=al&searchId=72B2D67C88B703BB6B7E888764A26CFC.kqQeWtawXauCyC8ghhRGJg%3A280&sid=72B2D67C88B703BB6B7E888764A26CFC.kqQeWtawXauCyC8ghhRGJg%3A280




#http://api.elsevier.com/content/search/index:SCOPUS?apiKey=392bff629e6d07593968592bbb629b0f,query=(AUTH-LAST-NAME(Chung) AND AUTH--FIRST(kevin) AND AFFIL(University of Michigan))


#unable to access service from my computer IP. WIll most likely need to develop and run from the IP that the licenseKey is registered to.

#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os, sys
import json
import requests
from pprint import pprint as pp
import arrow
from itertools import groupby
from stripper import strip_tags
import codecs
import requests
from docx import Document

def get_abstract_info(EID, refresh=False):
    'Get and save the json data for EID.'

    base = 'scopus-data/get_abstract_info'
    if not os.path.exists(base):
        os.makedirs(base)

    # Got it cached?  Return it.
    fname = '{0}/{1}'.format(base, EID)
    if os.path.exists(fname) and not refresh:
        with codecs.open(fname, "r", "utf-8") as f:
            return json.loads(f.read())

    # Otherwise retrieve and save results
    url = ("http://api.elsevier.com/content/abstract/eid/" + EID)
    resp = requests.get(url,
                        headers={'Accept':'application/json',
                                 'X-ELS-APIKey': MY_API_KEY})
    results = json.loads(resp.text)

    # Put it in the cache
    with codecs.open(fname, "w", "utf-8") as f:
        f.write(json.dumps(results))

    # Return it
    return results

def get_author_link(EID):

    data = get_abstract_info(EID)
    result = data['abstracts-retrieval-response']

    s = '<a href="http://www.scopus.com/authid/detail.url?origin=AuthorProfile&authorId={0}">{1}</a>'

    if 'authors' in result.keys():
        authors = [s.format(auid, name) for auid, name in
            zip([x['@auid'] for x in result['authors']['author']],
               [x['ce:indexed-name'] for x in result['authors']['author']])]
    else:
        authors = ""

# only gets first author...
#               zip([x['@auid'] for x in result['coredata']['dc:creator']['author']],
#                   [x['ce:indexed-name'] for x in result['coredata']['dc:creator']['author']])]

# yet another location to look for authors...
#               zip([x['@auid'] for x in result['bibrecord']['head']['author-group']['author']],
#                   [x['ce:indexed-name'] for x in result['bibrecord']['head']['author-group']['author']])]

    return ','.join(authors)

def get_journal_link(EID):

    data = get_abstract_info(EID)
    result = data['abstracts-retrieval-response']
    if 'source-id' in result['coredata'].keys():
        sid = result['coredata']['source-id']
    else:
        sid = ""
    if 'prism:publicationName' in result['coredata'].keys():
        journal = result['coredata']['prism:publicationName']
    else:
        journal = ""

    s = '<a href="http://www.scopus.com/source/sourceInfo.url?sourceId={sid}">{journal}</a>.'

    return s.format(sid=sid, journal=journal)

def get_doi_link(EID):

    data = get_abstract_info(EID)
    result = data['abstracts-retrieval-response']

    s = '<a href="http://dx.doi.org/{doi}">doi:{doi}</a>.'

    if 'prism:doi' in result['coredata']:
        return s.format(doi=result['coredata']['prism:doi'])
    else:
        return ""

# not used - ridiculous
def get_cite_img_link(EID):
    data = get_abstract_info(EID)
    result = data['abstracts-retrieval-response']
    s = '<img src="http://api.elsevier.com/content/abstract/citation-count?doi={doi}&httpAccept=image/jpeg&apiKey={apikey}"></img>'
    if 'prism:doi' in result['coredata']:
        return s.format(doi=result['coredata']['prism:doi'].strip(), apikey=MY_API_KEY)
    else:
        return ""

def get_abstract_link(EID):

    data = get_abstract_info(EID)
    result = data['abstracts-retrieval-response']
    title = result['coredata']['dc:title'].encode('utf-8')
    for ref in result['coredata']['link']:
        if ref['@rel'] == 'scopus':
            link = ref['@href'].replace('&amp;', '&').encode('utf-8')

    s = u'<a href="{link}">{title}</a>'.encode('utf-8')
    return s.format(link=link, title=title)

def get_html(EID):

    data = get_abstract_info(EID)

    result = data['abstracts-retrieval-response']

    s = '<li>{authors}, <i>{title}</i>, {journal} {coverDate};<b>{volume}</b>{pages} {doi}</li>'

    # keep it tidy - no Nones are wanted!
    issue = ''
    if result['coredata'].get('prism:issue'):
        issue = '({})'.format(result['coredata'].get('prism:issue'))
    volume = ''
    if result['coredata'].get('prism:volume'):
        volume = '{}'.format(result['coredata'].get('prism:volume'))
    # Combined format if both exist: volume(issue)
    if issue:
        volume = ("%s%s" % (volume,issue))

    pages = ''
    if result['coredata'].get('prism:pageRange'):
        pages = ':{}.'.format(result['coredata'].get('prism:pageRange'))

    doi = '' # TODO check exists before wrapping it in tags.
    doi = get_doi_link(EID)

    return s.format(authors=get_author_link(EID),
                    title=get_abstract_link(EID),
                    journal=get_journal_link(EID),
                    volume=volume,
                    pages=pages,
                    coverDate=result['coredata'].get('prism:coverDate'),
                    doi=doi,
                    cites=get_cite_img_link(EID))

if __name__ == '__main__':
    if len(sys.argv) < 3:
        sys.exit('Usage: %s <author_id> (html|htmlgroup|word)' % sys.argv[0])

    # very lame arg parsing
    auid = sys.argv[1]
    format = sys.argv[2]

    if not format:
        format = 'html'

    if not auid:
        # Filler for testing
        auid = "7005402670" # Alec Gallimore
        #auid = "7006840638" # David C. Munson
        #auid = "7410343879" # Stewart C. Wang

    # base for author loookup call
    api_url = "http://api.elsevier.com/content/search/scopus"

    # flags and options to try
    #view = "COMPLETE"
    view = "STANDARD"
    suppress = "true" # attempt to trim cruft from response
    count = "200" # maximum allowed for this service - api limit
    # TODO check total in response and page thru bibentries

    # our call
    s = "{api_url}?query=au-id({auid})&view={view}&suppressNavLinks={suppress}&start=0&count={count}"
    url = s.format( api_url = api_url,
                    auid = auid,
                    view = view,
                    suppress = suppress,
                    count = count)

    resp = requests.get(url,
                        headers={'Accept':'application/json',
                                 'X-ELS-APIKey': MY_API_KEY})

    author_data = json.loads(resp.text)
    entries = author_data['search-results']['entry']

    # Double-check date-descending sorting,
    # Tweak each entry with handy extra keys
    for entry in entries:
        publicationYear = arrow.get(entry['prism:coverDate']).year
        entry['pubyear'] = publicationYear
        entry['trim_eid'] = str(entry['eid']).rsplit(":",1)[-1]
    sorted_entries = sorted(entries, key=lambda k: k['prism:coverDate'], reverse=True)

    # Only three ways out of here
    if format == "word":
        doc = Document()
        for idx, entry in enumerate(sorted_entries, start=1):
            bib_entry = u''
            bib_entry = get_html(entry['trim_eid']).decode("utf-8")
            run = doc.add_paragraph().add_run()
            # Strip tags.  no formatting remains
            stripped = strip_tags(bib_entry)
            # Number that list - could likely do in Docx as native list
            numbered = "%s.  %s" % (idx, stripped)
            run.add_text(numbered)
        doc.save("bibliography.docx")
        # now just open "bibliography.docx" in Word

    elif format == "htmlgroup":
        print "<html lang=\"en-US\"><head><meta http-equiv=\"Content-Type\" content=\"text/html; charset=utf-8\" /><style>h1 {font-size: 1.5em} li {padding-bottom: .5em}</style></head><body><ol>"
        for year, pubs in groupby(sorted_entries, lambda x: x['pubyear']):
            print ("<h1>%s</h1>" % year)
            for pub in pubs:
                bib_entry = get_html(pub['trim_eid'])
                print bib_entry
        print "</ol></body></html>"
        # pipe to .html file and open in browser

    else: # html
        print "<html lang=\"en-US\"><head><meta http-equiv=\"Content-Type\" content=\"text/html; charset=utf-8\" /><style>h1 {font-size: 1.5em} li {padding-bottom: .5em}</style></head><body><ol>"
        for year, pubs in groupby(sorted_entries, lambda x: x['pubyear']):
            for pub in pubs:
                bib_entry = get_html(pub['trim_eid'])
                print bib_entry
        print "</ol></body></html>"
        # pipe to .html file and open in browser

